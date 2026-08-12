"""
Unit tests for document_loader. Uses real temporary files (txt and docx)
so we're testing actual parsing, not just mocked behavior.

    pytest backend/tests/test_document_loader.py -v
"""

import pytest
from docx import Document as DocxDocument

from backend.rag.document_loader import load_document


def test_load_txt(tmp_path):
    file_path = tmp_path / "sample.txt"
    file_path.write_text("Hello, this is a test document.\nSecond line.")

    doc = load_document(file_path)

    assert doc.filename == "sample.txt"
    assert len(doc.pages) == 1
    assert "Hello, this is a test document." in doc.full_text


def test_load_docx(tmp_path):
    file_path = tmp_path / "sample.docx"
    docx_doc = DocxDocument()
    docx_doc.add_paragraph("This is a real docx paragraph.")
    docx_doc.add_paragraph("This is a second paragraph.")
    docx_doc.save(file_path)

    doc = load_document(file_path)

    assert doc.filename == "sample.docx"
    assert "real docx paragraph" in doc.full_text
    assert "second paragraph" in doc.full_text


def test_unsupported_file_type_raises_clear_error(tmp_path):
    file_path = tmp_path / "sample.csv"
    file_path.write_text("a,b,c\n1,2,3")

    with pytest.raises(ValueError, match="Unsupported file type"):
        load_document(file_path)


def test_empty_txt_file_produces_no_pages(tmp_path):
    file_path = tmp_path / "empty.txt"
    file_path.write_text("")

    doc = load_document(file_path)

    # An empty file should load without crashing -- ingest.py is
    # responsible for deciding what to do with zero pages, not the loader.
    assert doc.full_text == ""
