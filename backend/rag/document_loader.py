"""
Document loader: turns a raw file (PDF, DOCX, TXT) into plain text.

Key principle: EXACT extraction, no summarizing, no interpretation.
Whatever the file says, that's what comes out -- word for word. Any
hallucination-prevention we do later only works if the text going IN
is a faithful copy of the source document.

We also keep track of *where* each piece of text came from (page number
for PDFs), because that becomes the citation later.
"""

from dataclasses import dataclass
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument


@dataclass
class PageText:
    """One page (or page-like unit) of extracted text, with its source location."""
    text: str
    page_number: int


@dataclass
class LoadedDocument:
    """A fully loaded document: its filename and a list of pages of text."""
    filename: str
    pages: list[PageText]

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)


def load_pdf(path: Path) -> LoadedDocument:
    pages = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            # x_tolerance/y_tolerance control how close two letters need to be
            # before pdfplumber treats them as "the same word" vs separate words.
            # The library default (x_tolerance=3) is too loose for some fonts
            # (e.g. many resume templates), causing words to run together with
            # no space -- "SoftwareEngineerwith" instead of "Software Engineer with".
            # Tightening it fixes this without hurting normal documents.
            text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            if text.strip():
                pages.append(PageText(text=text, page_number=i))
    return LoadedDocument(filename=path.name, pages=pages)


def load_docx(path: Path) -> LoadedDocument:
    doc = DocxDocument(path)
    # DOCX has no native "pages" concept, so we treat the whole doc as page 1.
    # (Good enough for citations at the document level; can be refined later.)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    return LoadedDocument(filename=path.name, pages=[PageText(text=text, page_number=1)])


def load_txt(path: Path) -> LoadedDocument:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return LoadedDocument(filename=path.name, pages=[PageText(text=text, page_number=1)])


def load_document(path: Path) -> LoadedDocument:
    """Dispatches to the right loader based on file extension."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(path)
    elif suffix == ".docx":
        return load_docx(path)
    elif suffix == ".txt":
        return load_txt(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
